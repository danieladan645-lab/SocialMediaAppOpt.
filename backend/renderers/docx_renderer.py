from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from models import AuditData

CORAL  = RGBColor(232, 96,  76)
TEAL   = RGBColor(61,  191, 191)
GOLD   = RGBColor(245, 166, 35)
DARK   = RGBColor(44,  44,  44)
GRAY   = RGBColor(150, 150, 150)
WHITE  = RGBColor(255, 255, 255)

CORAL_HEX  = "E8604C"
TEAL_HEX   = "3DBFBF"
CORAL_LIGHT_HEX = "FCE6E2"
TEAL_LIGHT_HEX  = "DCF5F5"
DARK_HEX   = "2C2C2C"


# === HELPERS ===

def _set_styles(doc):
    defs = {
        "Heading 1": (22, True, CORAL),
        "Heading 2": (16, True, TEAL),
        "Heading 3": (13, True, DARK),
    }
    for name, (size, bold, color) in defs.items():
        try:
            style = doc.styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = color
        except Exception:
            pass


def _para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix + ": ")
        rb.bold = True
        rb.font.name = "Calibri"
        rb.font.size = Pt(11)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _table(doc, headers, rows, header_hex=TEAL_HEX, alt_hex=TEAL_LIGHT_HEX):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        _shade_cell(cell, header_hex)
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        hex_fill = alt_hex if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            _shade_cell(cell, hex_fill)
    return tbl


def _footer(doc, handle, date):
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].clear()
    run = footer.paragraphs[0].add_run(
        f"Confidential  |  {handle} Instagram Audit  |  {date}"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


# === SECTION BUILDERS ===

def _cover(doc, d: AuditData):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Instagram Profile Audit Report")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = CORAL

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{d.handle} — {d.display_name}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(16)
    r2.font.color.rgb = TEAL

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"{d.date}  |  Instagram Profile Audit  |  {d.tier}")
    r3.font.name = "Calibri"
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY
    r3.italic = True

    doc.add_paragraph()
    doc.add_page_break()


def _section_executive(doc, d: AuditData):
    doc.add_heading("1. Executive Summary", level=1)
    _para(doc, (
        f"This audit evaluates the Instagram profile of {d.handle} ({d.display_name}). "
        f"The profile is assessed across six key dimensions: bio and CTA effectiveness, "
        "visual grid consistency, Highlight organization, content strategy, persona alignment, "
        "and conversion readiness."
    ))
    doc.add_paragraph()
    _para(doc, f"Overall Score: {d.overall_score}", bold=True, size=14, color=GOLD)
    doc.add_paragraph()
    doc.add_heading("Quick Wins", level=3)
    for qw in d.quick_wins:
        _bullet(doc, qw)
    doc.add_paragraph()


def _section_snapshot(doc, d: AuditData):
    doc.add_heading("2. Profile Snapshot", level=1)
    snap = d.profile_snapshot
    rows = [
        ("Username", snap.username),
        ("Display Name", snap.display_name),
        ("Followers", snap.followers),
        ("Bio Link", snap.bio_link),
        ("Profile Photo", snap.photo_desc),
        ("Platforms", snap.platforms),
    ]
    _table(doc, ["Attribute", "Detail"], rows)
    doc.add_paragraph()


def _section_bio(doc, d: AuditData):
    doc.add_heading("3. Bio Analysis", level=1)
    for block in d.bio_blocks:
        doc.add_heading(f"{block.label} — Score: {block.score_str}", level=3)
        _para(doc, block.body)
    doc.add_paragraph()


def _section_grid(doc, d: AuditData):
    doc.add_heading(f"4. Visual Grid & Aesthetic — Score: {d.grid_score}", level=1)
    for b in d.grid_bullets:
        _bullet(doc, b)
    doc.add_paragraph()


def _section_highlights(doc, d: AuditData):
    doc.add_heading(f"5. Highlights Analysis — Score: {d.highlight_score}", level=1)
    rows = [(r.name, r.cover_quality, r.recommendation) for r in d.highlight_rows]
    _table(doc, ["Highlight Name", "Cover Quality", "Recommendation"], rows)
    doc.add_paragraph()
    _para(doc, d.highlight_action, italic=True)
    doc.add_paragraph()


def _section_content(doc, d: AuditData):
    doc.add_heading(f"6. Content Strategy — Score: {d.content_score}", level=1)
    rows = [(b.label, f"{b.pct}%") for b in d.content_bars]
    _table(doc, ["Content Type", "Estimated % of Grid"], rows)
    doc.add_paragraph()
    for b in d.content_bullets:
        _bullet(doc, b)
    doc.add_paragraph()


def _section_persona(doc, persona, number: int):
    doc.add_heading(f"{number}. Customer Persona: {persona.title}", level=2)
    rows = [
        ("Name & Age", persona.name),
        ("Location", persona.location),
        ("Goals", persona.goals),
        ("Pain Points", persona.pain_points),
        ("How They Found Profile", persona.found_via),
        ("First Impression", persona.first_impression),
        ("What They Need", persona.needs),
        ("Conversion Path", persona.conversion),
    ]
    _table(doc, ["Field", "Detail"], rows)
    doc.add_paragraph()
    _para(doc, "Strategic Implication: " + persona.implication, italic=True)
    doc.add_paragraph()


def _section_strengths(doc, d: AuditData, section_num: int):
    doc.add_heading(f"{section_num}. Strengths", level=1)
    for s in d.strengths:
        _bullet(doc, s.body, bold_prefix=s.title)
    doc.add_paragraph()


def _section_weaknesses(doc, d: AuditData, section_num: int):
    doc.add_heading(f"{section_num}. Weaknesses & Opportunities", level=1)
    rows = list(zip(d.weaknesses, d.opportunities))
    _table(doc, ["Weaknesses", "Opportunities"], rows,
           header_hex=CORAL_HEX, alt_hex=CORAL_LIGHT_HEX)
    doc.add_paragraph()


def _section_recommendations(doc, d: AuditData, section_num: int):
    doc.add_heading(f"{section_num}. Recommendations & Action Plan", level=1)
    rows = [(r.priority, r.action, r.impact) for r in d.recommendations]
    _table(doc, ["Priority", "Action", "Expected Impact"], rows,
           header_hex=DARK_HEX, alt_hex=CORAL_LIGHT_HEX)
    doc.add_paragraph()


def _section_scorecard(doc, d: AuditData, section_num: int):
    doc.add_heading(f"{section_num}. Score Summary", level=1)
    rows = [(r.category, r.score_str) for r in d.score_rows]
    _table(doc, ["Category", "Score"], rows)
    doc.add_paragraph()
    _para(doc, f"Audit conducted: {d.date}  |  {d.handle} Instagram Profile",
          italic=True, color=GRAY)


# === ENTRY POINT ===

def build_docx(d: AuditData) -> BytesIO:
    doc = Document()
    _set_styles(doc)
    _footer(doc, d.handle, d.date)

    _cover(doc, d)
    _section_executive(doc, d)
    _section_snapshot(doc, d)
    _section_bio(doc, d)
    _section_grid(doc, d)
    _section_highlights(doc, d)
    _section_content(doc, d)

    persona_section_start = 7
    for i, persona in enumerate(d.personas):
        _section_persona(doc, persona, persona_section_start + i)

    offset = persona_section_start + len(d.personas)
    _section_strengths(doc, d, offset)
    _section_weaknesses(doc, d, offset + 1)
    _section_recommendations(doc, d, offset + 2)
    _section_scorecard(doc, d, offset + 3)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
